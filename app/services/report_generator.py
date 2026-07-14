# -*- coding: utf-8 -*-
"""
报表自动生成服务
负责从数据库聚合报表数据并生成 Word 文档
支持：舆情日报、舆情周报、事件专报
"""
import os
import json
from datetime import datetime, timedelta
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn

from app.core.database import get_connection

# 报表文件输出目录
REPORTS_DIR = Path(__file__).resolve().parent.parent.parent / "data" / "reports"


def _ensure_reports_dir():
    """确保报表输出目录存在"""
    REPORTS_DIR.mkdir(parents=True, exist_ok=True)


def _today_str():
    """获取当日日期字符串 YYYY-MM-DD"""
    return datetime.now().strftime("%Y-%m-%d")


def _now_str():
    """获取当前时间字符串"""
    return datetime.now().strftime("%Y-%m-%d %H:%M:%S")


def _rows_to_dicts(rows):
    """将 sqlite3.Row 转为字典列表"""
    return [dict(r) for r in rows]


def _safe_json(text, default=None):
    """安全解析 JSON 字符串"""
    if default is None:
        default = []
    try:
        return json.loads(text) if text else default
    except (json.JSONDecodeError, TypeError):
        return default


def _set_cell_font(cell, text, bold=False, size=9):
    """设置表格单元格的字体"""
    cell.text = ""
    run = cell.paragraphs[0].add_run(str(text))
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def _add_heading(doc, text, level=1):
    """添加标题段落"""
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    return heading


def _add_paragraph(doc, text, bold=False, size=10):
    """添加正文段落"""
    para = doc.add_paragraph()
    run = para.add_run(str(text))
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    return para


def _add_placeholder(doc, desc):
    """添加图表占位标注"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(f"[图表占位：{desc}]")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(150, 150, 150)
    run.font.italic = True
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    return para


def _create_table(doc, headers, rows):
    """创建通用表格并填充数据"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, header in enumerate(headers):
        _set_cell_font(table.rows[0].cells[i], header, bold=True, size=9)
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            _set_cell_font(table.rows[r_idx + 1].cells[c_idx], str(val), size=9)
    return table


# ============================================================
# 数据聚合函数
# ============================================================

def gather_daily_report_data() -> dict:
    """
    聚合当日日报数据
    返回包含以下内容的字典：
    - 当日新增热点事件数
    - 当日事件热度 TOP10
    - 各事件情感分布汇总
    - 当日高风险事件列表
    - 当日虚假文本标记统计
    - 各平台新闻数量
    """
    today = _today_str()
    conn = get_connection()
    try:
        # 1. 当日新增热点事件数
        count_row = conn.execute(
            "SELECT COUNT(*) AS cnt FROM hot_event WHERE DATE(created_at) = ?",
            (today,)
        ).fetchone()
        new_event_count = count_row["cnt"] if count_row else 0

        # 2. 当日事件热度 TOP10
        top10_events = conn.execute(
            "SELECT * FROM hot_event WHERE DATE(created_at) = ? ORDER BY heat_score DESC LIMIT 10",
            (today,)
        ).fetchall()
        top10_events = _rows_to_dicts(top10_events)

        # 3. 各事件情感分布汇总（当日有分析记录的事件）
        sentiment_rows = conn.execute(
            """
            SELECT ea.event_id, he.title, ea.positive_ratio, ea.negative_ratio, ea.neutral_ratio
            FROM event_analysis ea
            JOIN hot_event he ON ea.event_id = he.id
            WHERE DATE(ea.analyzed_at) = ?
            ORDER BY ea.event_id
            """,
            (today,)
        ).fetchall()
        sentiment_summary = _rows_to_dicts(sentiment_rows)

        # 4. 当日高风险事件列表（risk_level=high/critical 或 credibility < 0.5）
        high_risk_events = conn.execute(
            "SELECT * FROM hot_event WHERE DATE(created_at) = ? AND (risk_level = 'high' OR risk_level = 'critical')",
            (today,)
        ).fetchall()
        high_risk_events = _rows_to_dicts(high_risk_events)

        # 从 event_analysis 中查找低可信度事件
        low_cred_events = conn.execute(
            """
            SELECT ea.event_id, he.title, ea.credibility_score
            FROM event_analysis ea
            JOIN hot_event he ON ea.event_id = he.id
            WHERE DATE(ea.analyzed_at) = ? AND ea.credibility_score < 0.5
            """,
            (today,)
        ).fetchall()
        low_cred_events = _rows_to_dicts(low_cred_events)

        # 5. 当日虚假文本标记统计（从 event_analysis 取 fake_flags）
        fake_flag_rows = conn.execute(
            """
            SELECT ea.event_id, he.title, ea.fake_flags
            FROM event_analysis ea
            JOIN hot_event he ON ea.event_id = he.id
            WHERE DATE(ea.analyzed_at) = ? AND ea.fake_flags != '[]' AND ea.fake_flags IS NOT NULL
            """,
            (today,)
        ).fetchall()
        fake_flag_events = _rows_to_dicts(fake_flag_rows)
        fake_flag_count = len(fake_flag_events)

        # 6. 各平台新闻数量（从 raw_news 按日期统计）
        platform_rows = conn.execute(
            """
            SELECT source_platform, COUNT(*) AS cnt
            FROM raw_news
            WHERE DATE(crawled_at) = ?
            GROUP BY source_platform
            ORDER BY cnt DESC
            """,
            (today,)
        ).fetchall()
        platform_news_stats = _rows_to_dicts(platform_rows)

        return {
            "report_date": today,
            "generated_at": _now_str(),
            "new_event_count": new_event_count,
            "top10_events": top10_events,
            "sentiment_summary": sentiment_summary,
            "high_risk_events": high_risk_events,
            "low_credibility_events": low_cred_events,
            "fake_flag_events": fake_flag_events,
            "fake_flag_count": fake_flag_count,
            "platform_news_stats": platform_news_stats,
        }
    finally:
        conn.close()


def gather_weekly_report_data() -> dict:
    """
    聚合本周数据（近7天）
    返回包含以下内容的字典：
    - 本周热度趋势（按天统计 hot_event 数量）
    - 周度情感变化（汇总本周 event_analysis 的情感比例）
    - 高频负面话题 TOP5
    - 虚假信息汇总
    - 重点传播节点统计
    """
    today = datetime.now()
    week_start = (today - timedelta(days=6)).strftime("%Y-%m-%d")
    week_end = today.strftime("%Y-%m-%d")
    conn = get_connection()
    try:
        # 1. 本周热度趋势（按天统计 hot_event 数量）
        trend_rows = conn.execute(
            """
            SELECT DATE(created_at) AS day, COUNT(*) AS event_count
            FROM hot_event
            WHERE DATE(created_at) >= ? AND DATE(created_at) <= ?
            GROUP BY DATE(created_at)
            ORDER BY day
            """,
            (week_start, week_end)
        ).fetchall()
        heat_trend = _rows_to_dicts(trend_rows)

        # 2. 周度情感变化（汇总本周 event_analysis 的情感比例）
        sentiment_rows = conn.execute(
            """
            SELECT
                AVG(ea.positive_ratio) AS avg_positive,
                AVG(ea.negative_ratio) AS avg_negative,
                AVG(ea.neutral_ratio) AS avg_neutral
            FROM event_analysis ea
            WHERE DATE(ea.analyzed_at) >= ? AND DATE(ea.analyzed_at) <= ?
            """,
            (week_start, week_end)
        ).fetchone()
        weekly_sentiment = dict(sentiment_rows) if sentiment_rows else {
            "avg_positive": 0, "avg_negative": 0, "avg_neutral": 0
        }

        # 3. 高频负面话题 TOP5（negative_ratio 最高的事件）
        negative_rows = conn.execute(
            """
            SELECT ea.event_id, he.title, ea.negative_ratio
            FROM event_analysis ea
            JOIN hot_event he ON ea.event_id = he.id
            WHERE DATE(ea.analyzed_at) >= ? AND DATE(ea.analyzed_at) <= ?
            ORDER BY ea.negative_ratio DESC
            LIMIT 5
            """,
            (week_start, week_end)
        ).fetchall()
        top_negative_topics = _rows_to_dicts(negative_rows)

        # 4. 虚假信息汇总（有 fake_flags 的事件列表）
        fake_rows = conn.execute(
            """
            SELECT ea.event_id, he.title, ea.fake_flags, ea.analyzed_at
            FROM event_analysis ea
            JOIN hot_event he ON ea.event_id = he.id
            WHERE DATE(ea.analyzed_at) >= ? AND DATE(ea.analyzed_at) <= ?
              AND ea.fake_flags != '[]' AND ea.fake_flags IS NOT NULL
            """,
            (week_start, week_end)
        ).fetchall()
        fake_info_summary = _rows_to_dicts(fake_rows)

        # 5. 重点传播节点统计（从 spread_info 取 graph_data.nodes）
        spread_rows = conn.execute(
            """
            SELECT si.event_id, he.title, si.graph_data, si.spread_depth,
                   si.total_reposts, si.total_reads
            FROM spread_info si
            JOIN hot_event he ON si.event_id = he.id
            WHERE DATE(si.traced_at) >= ? AND DATE(si.traced_at) <= ?
            """,
            (week_start, week_end)
        ).fetchall()
        spread_data = _rows_to_dicts(spread_rows)

        # 提取传播节点统计
        key_spread_nodes = []
        for item in spread_data:
            graph_data = _safe_json(item.get("graph_data", "{}"), default={})
            nodes = graph_data.get("nodes", [])
            if nodes:
                key_spread_nodes.extend(nodes)

        return {
            "week_start": week_start,
            "week_end": week_end,
            "generated_at": _now_str(),
            "heat_trend": heat_trend,
            "weekly_sentiment": weekly_sentiment,
            "top_negative_topics": top_negative_topics,
            "fake_info_summary": fake_info_summary,
            "spread_data": spread_data,
            "key_spread_nodes": key_spread_nodes,
        }
    finally:
        conn.close()


def gather_event_report_data(event_id: int) -> dict:
    """
    针对单个事件的数据聚合
    返回包含以下内容的字典：
    - 事件基本信息
    - 事件分析数据（情感分布、关键词、平台覆盖、可信度）
    - 传播溯源数据（spread_nodes, graph_data）
    - 关联新闻列表
    """
    conn = get_connection()
    try:
        # 1. 事件基本信息
        event_row = conn.execute(
            "SELECT * FROM hot_event WHERE id = ?", (event_id,)
        ).fetchone()
        event_info = dict(event_row) if event_row else None

        if event_info is None:
            return {"error": f"事件 ID {event_id} 不存在"}

        # 2. 事件分析数据
        analysis_row = conn.execute(
            "SELECT * FROM event_analysis WHERE event_id = ? ORDER BY analyzed_at DESC LIMIT 1",
            (event_id,)
        ).fetchone()
        analysis_data = dict(analysis_row) if analysis_row else None

        if analysis_data:
            # 解析 JSON 字段
            analysis_data["high_freq_keywords_parsed"] = _safe_json(analysis_data.get("high_freq_keywords"))
            analysis_data["platform_coverage_parsed"] = _safe_json(analysis_data.get("platform_coverage"), default={})
            analysis_data["fake_flags_parsed"] = _safe_json(analysis_data.get("fake_flags"))

        # 3. 传播溯源数据
        spread_row = conn.execute(
            "SELECT * FROM spread_info WHERE event_id = ? ORDER BY traced_at DESC LIMIT 1",
            (event_id,)
        ).fetchone()
        spread_data = dict(spread_row) if spread_row else None

        if spread_data:
            spread_data["spread_nodes_parsed"] = _safe_json(spread_data.get("spread_nodes"))
            spread_data["graph_data_parsed"] = _safe_json(spread_data.get("graph_data"), default={})

        # 4. 关联新闻列表（通过 raw_news 的 event_id 或通过标题模糊匹配）
        related_news = conn.execute(
            "SELECT * FROM raw_news WHERE event_id = ? ORDER BY crawled_at DESC",
            (event_id,)
        ).fetchall()
        related_news = _rows_to_dicts(related_news)

        return {
            "event_info": event_info,
            "analysis_data": analysis_data,
            "spread_data": spread_data,
            "related_news": related_news,
            "generated_at": _now_str(),
        }
    finally:
        conn.close()


# ============================================================
# Word 文档生成
# ============================================================

def generate_report_docx(data: dict, report_type: str) -> str:
    """
    根据聚合数据生成 Word 文档

    参数:
        data: 聚合后的数据字典
        report_type: 报表类型 "daily" / "weekly" / "event"

    返回:
        生成的 Word 文件绝对路径
    """
    _ensure_reports_dir()

    doc = Document()

    # 设置默认字体
    style = doc.styles["Normal"]
    font = style.font
    font.name = "微软雅黑"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    font.size = Pt(10)

    # ---- 根据类型生成不同内容 ----
    if report_type == "daily":
        _build_daily_doc(doc, data)
        filename = f"舆情日报_{data.get('report_date', _today_str())}.docx"
    elif report_type == "weekly":
        _build_weekly_doc(doc, data)
        filename = f"舆情周报_{data.get('week_start', '')}_{data.get('week_end', '')}.docx"
    elif report_type == "event":
        event_info = data.get("event_info", {})
        event_name = event_info.get("title", "未知事件") if event_info else "未知事件"
        _build_event_doc(doc, data)
        filename = f"事件专报_{event_name}.docx"
    else:
        raise ValueError(f"不支持的报表类型: {report_type}")

    # 清理文件名中的非法字符
    filename = "".join(c if c not in r'/\\:*?"<>|' else "_" for c in filename)

    filepath = str(REPORTS_DIR / filename)
    doc.save(filepath)
    return filepath


def _build_daily_doc(doc: Document, data: dict):
    """构建舆情日报 Word 文档内容"""
    report_date = data.get("report_date", "")
    generated_at = data.get("generated_at", "")

    # 标题
    title = doc.add_heading("舆情日报", level=0)
    for run in title.runs:
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    _add_paragraph(doc, f"报告日期：{report_date}")
    _add_paragraph(doc, f"生成时间：{generated_at}")

    doc.add_paragraph()  # 空行分隔

    # ---- 模块一：当日数据概览 ----
    _add_heading(doc, "一、当日数据概览", level=1)
    _add_paragraph(doc, f"当日新增热点事件数：{data.get('new_event_count', 0)}")

    # ---- 模块二：事件热度 TOP10 ----
    _add_heading(doc, "二、事件热度 TOP10", level=1)
    top10 = data.get("top10_events", [])
    if top10:
        headers = ["排名", "事件标题", "热度分数", "风险等级", "生命周期"]
        rows = []
        for i, evt in enumerate(top10, 1):
            rows.append([
                str(i),
                evt.get("title", ""),
                f"{evt.get('heat_score', 0):.1f}",
                _risk_label(evt.get("risk_level", "")),
                _lifecycle_label(evt.get("lifecycle", "")),
            ])
        _create_table(doc, headers, rows)
    else:
        _add_paragraph(doc, "当日暂无热点事件数据")
    _add_placeholder(doc, "热度 TOP10 柱状图")

    # ---- 模块三：情感分布汇总 ----
    _add_heading(doc, "三、情感分布汇总", level=1)
    sentiments = data.get("sentiment_summary", [])
    if sentiments:
        headers = ["事件标题", "正面占比", "负面占比", "中性占比"]
        rows = []
        for s in sentiments:
            rows.append([
                s.get("title", ""),
                f"{s.get('positive_ratio', 0) * 100:.1f}%",
                f"{s.get('negative_ratio', 0) * 100:.1f}%",
                f"{s.get('neutral_ratio', 0) * 100:.1f}%",
            ])
        _create_table(doc, headers, rows)
    else:
        _add_paragraph(doc, "当日暂无情感分析数据")
    _add_placeholder(doc, "情感分布饼图")

    # ---- 模块四：高风险事件 ----
    _add_heading(doc, "四、高风险事件", level=1)
    high_risk = data.get("high_risk_events", [])
    low_cred = data.get("low_credibility_events", [])
    if high_risk:
        _add_paragraph(doc, "高风险事件（风险等级为 high/critical）：", bold=True)
        headers = ["事件标题", "热度分数", "风险等级", "生命周期"]
        rows = []
        for evt in high_risk:
            rows.append([
                evt.get("title", ""),
                f"{evt.get('heat_score', 0):.1f}",
                evt.get("risk_level", ""),
                _lifecycle_label(evt.get("lifecycle", "")),
            ])
        _create_table(doc, headers, rows)
    else:
        _add_paragraph(doc, "当日无高风险事件")
    if low_cred:
        _add_paragraph(doc, "低可信度事件（可信度 < 0.5）：", bold=True)
        headers = ["事件标题", "可信度分数"]
        rows = [[e.get("title", ""), f"{e.get('credibility_score', 0):.2f}"] for e in low_cred]
        _create_table(doc, headers, rows)

    # ---- 模块五：虚假文本标记 ----
    _add_heading(doc, "五、虚假文本标记统计", level=1)
    _add_paragraph(doc, f"当日标记虚假文本事件数：{data.get('fake_flag_count', 0)}")
    fake_events = data.get("fake_flag_events", [])
    if fake_events:
        headers = ["事件标题", "虚假特征标记"]
        rows = []
        for fe in fake_events:
            flags = _safe_json(fe.get("fake_flags"))
            rows.append([fe.get("title", ""), "、".join(flags) if flags else "无"])
        _create_table(doc, headers, rows)

    # ---- 模块六：各平台新闻数量 ----
    _add_heading(doc, "六、各平台新闻数量", level=1)
    platform_stats = data.get("platform_news_stats", [])
    if platform_stats:
        headers = ["平台", "新闻数量"]
        rows = [[p.get("source_platform", ""), str(p.get("cnt", 0))] for p in platform_stats]
        _create_table(doc, headers, rows)
    else:
        _add_paragraph(doc, "当日暂无平台新闻数据")
    _add_placeholder(doc, "各平台新闻数量分布图")


def _build_weekly_doc(doc: Document, data: dict):
    """构建舆情周报 Word 文档内容"""
    week_start = data.get("week_start", "")
    week_end = data.get("week_end", "")
    generated_at = data.get("generated_at", "")

    # 标题
    title = doc.add_heading("舆情周报", level=0)
    for run in title.runs:
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    _add_paragraph(doc, f"统计周期：{week_start} 至 {week_end}")
    _add_paragraph(doc, f"生成时间：{generated_at}")

    doc.add_paragraph()

    # ---- 模块一：热度趋势 ----
    _add_heading(doc, "一、本周热度趋势", level=1)
    heat_trend = data.get("heat_trend", [])
    if heat_trend:
        headers = ["日期", "新增事件数"]
        rows = [[t.get("day", ""), str(t.get("event_count", 0))] for t in heat_trend]
        _create_table(doc, headers, rows)
    else:
        _add_paragraph(doc, "本周暂无热度趋势数据")
    _add_placeholder(doc, "热度趋势折线图")

    # ---- 模块二：周度情感变化 ----
    _add_heading(doc, "二、周度情感变化", level=1)
    ws = data.get("weekly_sentiment", {})
    headers = ["情感类型", "平均占比"]
    rows = [
        ["正面", f"{ws.get('avg_positive', 0) * 100:.1f}%"],
        ["负面", f"{ws.get('avg_negative', 0) * 100:.1f}%"],
        ["中性", f"{ws.get('avg_neutral', 0) * 100:.1f}%"],
    ]
    _create_table(doc, headers, rows)
    _add_placeholder(doc, "周度情感变化趋势图")

    # ---- 模块三：高频负面话题 TOP5 ----
    _add_heading(doc, "三、高频负面话题 TOP5", level=1)
    top_neg = data.get("top_negative_topics", [])
    if top_neg:
        headers = ["排名", "事件标题", "负面占比"]
        rows = []
        for i, topic in enumerate(top_neg, 1):
            rows.append([
                str(i),
                topic.get("title", ""),
                f"{topic.get('negative_ratio', 0) * 100:.1f}%",
            ])
        _create_table(doc, headers, rows)
    else:
        _add_paragraph(doc, "本周暂无高频负面话题")
    _add_placeholder(doc, "负面话题 TOP5 柱状图")

    # ---- 模块四：虚假信息汇总 ----
    _add_heading(doc, "四、虚假信息汇总", level=1)
    fake_info = data.get("fake_info_summary", [])
    _add_paragraph(doc, f"本周涉及虚假信息事件数：{len(fake_info)}")
    if fake_info:
        headers = ["事件标题", "虚假特征标记", "分析日期"]
        rows = []
        for fi in fake_info:
            flags = _safe_json(fi.get("fake_flags"))
            rows.append([
                fi.get("title", ""),
                "、".join(flags) if flags else "无",
                fi.get("analyzed_at", ""),
            ])
        _create_table(doc, headers, rows)
    else:
        _add_paragraph(doc, "本周暂无虚假信息标记")

    # ---- 模块五：重点传播节点统计 ----
    _add_heading(doc, "五、重点传播节点统计", level=1)
    spread_data = data.get("spread_data", [])
    if spread_data:
        headers = ["事件标题", "传播深度", "转发总量", "阅读总量"]
        rows = []
        for sd in spread_data:
            rows.append([
                sd.get("title", ""),
                str(sd.get("spread_depth", 0)),
                str(sd.get("total_reposts", 0)),
                str(sd.get("total_reads", 0)),
            ])
        _create_table(doc, headers, rows)
        _add_placeholder(doc, "传播关系网络图")
    else:
        _add_paragraph(doc, "本周暂无传播溯源数据")

    key_nodes = data.get("key_spread_nodes", [])
    if key_nodes:
        _add_paragraph(doc, f"重点传播节点数：{len(key_nodes)}")


def _build_event_doc(doc: Document, data: dict):
    """构建事件专报 Word 文档内容"""
    if "error" in data:
        _add_paragraph(doc, data["error"])
        return

    event_info = data.get("event_info", {})
    analysis = data.get("analysis_data", {})
    spread = data.get("spread_data", {})
    related_news = data.get("related_news", [])
    generated_at = data.get("generated_at", "")

    # 标题
    event_name = event_info.get("title", "未知事件")
    title = doc.add_heading(f"事件专报：{event_name}", level=0)
    for run in title.runs:
        run.font.name = "微软雅黑"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    _add_paragraph(doc, f"生成时间：{generated_at}")
    doc.add_paragraph()

    # ---- 模块一：事件基本信息 ----
    _add_heading(doc, "一、事件基本信息", level=1)
    headers = ["字段", "内容"]
    rows = [
        ["事件标题", event_info.get("title", "")],
        ["事件概述", event_info.get("summary", "无")],
        ["热度分数", f"{event_info.get('heat_score', 0):.1f}"],
        ["风险等级", _risk_label(event_info.get("risk_level", ""))],
        ["生命周期", _lifecycle_label(event_info.get("lifecycle", ""))],
        ["创建时间", event_info.get("created_at", "")],
        ["更新时间", event_info.get("updated_at", "")],
    ]
    _create_table(doc, headers, rows)

    # ---- 模块二：事件分析数据 ----
    _add_heading(doc, "二、事件分析数据", level=1)
    if analysis:
        # 情感分布
        _add_paragraph(doc, "情感分布：", bold=True)
        headers = ["情感类型", "占比"]
        rows = [
            ["正面", f"{analysis.get('positive_ratio', 0) * 100:.1f}%"],
            ["负面", f"{analysis.get('negative_ratio', 0) * 100:.1f}%"],
            ["中性", f"{analysis.get('neutral_ratio', 0) * 100:.1f}%"],
        ]
        _create_table(doc, headers, rows)
        _add_placeholder(doc, "事件情感分布饼图")

        # 高频关键词
        keywords = analysis.get("high_freq_keywords_parsed", [])
        if keywords:
            _add_paragraph(doc, "高频关键词：", bold=True)
            _add_paragraph(doc, "、".join(keywords))

        # 平台覆盖
        coverage = analysis.get("platform_coverage_parsed", {})
        if coverage:
            _add_paragraph(doc, "平台覆盖：", bold=True)
            headers = ["平台", "报道占比"]
            rows = [[k, f"{v * 100:.1f}%"] for k, v in coverage.items()]
            _create_table(doc, headers, rows)

        # 可信度
        cred = analysis.get("credibility_score", 0)
        _add_paragraph(doc, f"可信度分数：{cred:.2f}")

        # 虚假特征标记
        fake_flags = analysis.get("fake_flags_parsed", [])
        if fake_flags:
            _add_paragraph(doc, "虚假特征标记：", bold=True)
            _add_paragraph(doc, "、".join(fake_flags))
        else:
            _add_paragraph(doc, "虚假特征标记：无")
    else:
        _add_paragraph(doc, "暂无分析数据")

    # ---- 模块三：传播溯源数据 ----
    _add_heading(doc, "三、传播溯源数据", level=1)
    if spread:
        _add_paragraph(doc, "首发平台：" + spread.get("origin_platform", "未知"))
        _add_paragraph(doc, "首发链接：" + (spread.get("origin_url") or "无"))

        headers = ["传播指标", "数据"]
        rows = [
            ["传播深度", str(spread.get("spread_depth", 0))],
            ["转发总量", str(spread.get("total_reposts", 0))],
            ["阅读总量", str(spread.get("total_reads", 0))],
        ]
        _create_table(doc, headers, rows)

        # 传播节点
        nodes = spread.get("spread_nodes_parsed", [])
        if nodes:
            _add_paragraph(doc, f"传播节点数：{len(nodes)}", bold=True)
            headers = ["节点名称", "平台"]
            rows = []
            for node in nodes[:20]:  # 最多展示20个节点
                name = node.get("name", node.get("id", "未知"))
                platform = node.get("platform", "未知")
                rows.append([str(name), str(platform)])
            _create_table(doc, headers, rows)
        _add_placeholder(doc, "传播关系网络图")
    else:
        _add_paragraph(doc, "暂无传播溯源数据")

    # ---- 模块四：关联新闻列表 ----
    _add_heading(doc, "四、关联新闻列表", level=1)
    if related_news:
        headers = ["序号", "新闻标题", "来源平台", "抓取时间"]
        rows = []
        for i, news in enumerate(related_news, 1):
            rows.append([
                str(i),
                news.get("title", ""),
                news.get("source_platform", ""),
                news.get("crawled_at", ""),
            ])
        _create_table(doc, headers, rows)
    else:
        _add_paragraph(doc, "暂无关联新闻数据")


def _risk_label(risk: str) -> str:
    """风险等级中文标签"""
    mapping = {
        "low": "低",
        "medium": "中",
        "high": "高",
        "critical": "严重",
    }
    return mapping.get(risk, risk)


def _lifecycle_label(lifecycle: str) -> str:
    """生命周期中文标签"""
    mapping = {
        "latent": "潜伏期",
        "growth": "增长期",
        "peak": "爆发期",
        "decline": "衰退期",
    }
    return mapping.get(lifecycle, lifecycle)
